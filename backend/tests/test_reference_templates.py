from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt

from app.review.parser import FileParser
from app.review.reference_templates import ReferenceTemplateEngine


def _docx(*, body: str, font: str = "Times New Roman", margin: float = 1.0) -> bytes:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = font
    normal.font.size = Pt(12)
    for section in doc.sections:
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(body)
    doc.add_heading("References", level=1)
    doc.add_paragraph("Example, 2026.")
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def test_learns_format_and_structure_without_learning_body_content() -> None:
    engine = ReferenceTemplateEngine()
    profile = engine.learn(_docx(body="Teacher placeholder text."), "school-template.docx")

    assert profile["content_policy"] == "format_and_structure_only"
    assert profile["body_text_is_user_owned"] is True
    assert profile["body"]["font_name"] == "Times New Roman"
    assert profile["required_sections"] == ["Introduction", "References"]
    assert "Teacher placeholder text." not in str(profile)


def test_compare_ignores_different_user_body_text() -> None:
    engine = ReferenceTemplateEngine()
    profile = engine.learn(_docx(body="Template placeholder."), "template.docx")
    student = FileParser().parse(
        _docx(body="Completely original student research content."),
        "student.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    issues = engine.compare(student, profile)
    assert issues == []


def test_compare_finds_format_mismatch() -> None:
    engine = ReferenceTemplateEngine()
    profile = engine.learn(_docx(body="Template."), "template.docx")
    student = FileParser().parse(
        _docx(body="Student.", font="Arial", margin=0.5),
        "student.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    issues = engine.compare(student, profile)
    rule_ids = {issue.rule_id for issue in issues}
    assert "template.body.font" in rule_ids
    assert "template.layout.margin_left_in" in rule_ids
    assert all(issue.source == "reference_template" for issue in issues)
    assert all(not issue.autofix_allowed for issue in issues)


def test_postgres_initialize_emits_valid_template_and_session_ddl(monkeypatch) -> None:
    from contextlib import contextmanager
    from app.database.postgres_adapter import PostgresAdapter

    statements: list[str] = []

    class Cursor:
        def execute(self, statement: str, *args) -> None:
            statements.append(statement.strip())
        def __enter__(self):
            return self
        def __exit__(self, exc_type, exc, tb):
            return False

    class Connection:
        committed = False
        def commit(self) -> None:
            self.committed = True

    connection = Connection()

    @contextmanager
    def connect():
        yield connection

    adapter = PostgresAdapter("postgresql://unused")
    monkeypatch.setattr(adapter, "_connect", connect)
    monkeypatch.setattr(adapter, "_cursor", lambda conn: Cursor())
    adapter.initialize()

    assert connection.committed is True
    assert any(sql.startswith("CREATE TABLE IF NOT EXISTS reference_templates") for sql in statements)
    assert any(sql.startswith("CREATE TABLE IF NOT EXISTS review_sessions") for sql in statements)
    assert all(not sql.startswith(")") for sql in statements)