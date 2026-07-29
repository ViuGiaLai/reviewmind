from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

from ..models import (
    BlockType,
    Citation,
    DocumentBlock,
    DocumentMetadata,
    DocumentModel,
    FigureData,
    Hyperlink,
    Sentence,
    SourceLocation,
    TableCell,
    TableData,
)
from .text import TextParser


class DocxParser:
    """Comprehensive DOCX parser with full structural extraction."""

    NS = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    }

    def __init__(self) -> None:
        self.text_parser = TextParser()

    # ── Main Entry ─────────────────────────────────────────────────────────────

    def parse(self, content: bytes, filename: str, content_type: str) -> DocumentModel:
        doc = DocxDocument(BytesIO(content))
        blocks: list[DocumentBlock] = []
        tables: list[TableData] = []
        figures: list[FigureData] = []
        footnotes: list[str] = []
        endnotes: list[str] = []
        revisions: list[dict[str, Any]] = []
        page_breaks: list[int] = []
        block_idx = 0

        # Global collections
        all_hyperlinks: list[Hyperlink] = []
        all_citations: list[Citation] = []
        all_sentences: list[Sentence] = []

        # ── Extract document metadata ────────────────────────────────────────
        metadata = self._extract_metadata(doc)

        # ── Extract headers/footers ──────────────────────────────────────────
        header_blocks, footer_blocks = self._extract_headers_footers(doc)
        blocks.extend(header_blocks)
        block_idx = len(blocks)

        # ── Extract main body paragraphs and tables ──────────────────────────
        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":  # Paragraph
                p_blocks, p_figures = self._parse_paragraph(element, doc, block_idx)
                for b in p_blocks:
                    from dataclasses import replace
                    if b.type == BlockType.PAGE_BREAK:
                        page_breaks.append(block_idx)
                    # Extract hyperlinks from paragraph (pass doc for relationship resolution)
                    b_hyperlinks, b_citations = self._extract_hyperlinks_and_citations(element, b, doc)
                    all_hyperlinks.extend(b_hyperlinks)
                    all_citations.extend(b_citations)
                    # Split sentences (use replace since DocumentBlock is frozen)
                    from ..evidence.engine import evidence_engine
                    sentences = evidence_engine.split_sentences(b.text, b.node_id)
                    all_sentences.extend(sentences)
                    b = replace(b, sentences=sentences, hyperlinks=b_hyperlinks, citations=b_citations)
                    blocks.append(b)
                    block_idx += 1
                figures.extend(p_figures)

            elif tag == "tbl":  # Table
                table_data, table_blocks = self._parse_table(element, doc, block_idx)
                if table_data:
                    tables.append(table_data)
                for b in table_blocks:
                    b_hyperlinks, b_citations = self._extract_hyperlinks_and_citations(element, b)
                    all_hyperlinks.extend(b_hyperlinks)
                    all_citations.extend(b_citations)
                blocks.extend(table_blocks)
                block_idx += len(table_blocks)

            elif tag == "sdt":  # Structured document tag (content controls)
                sdt_blocks, sdt_figures, sdt_tables = self._parse_sdt(element, doc, block_idx)
                blocks.extend(sdt_blocks)
                figures.extend(sdt_figures)
                tables.extend(sdt_tables)
                block_idx += len(sdt_blocks)

        # ── Extract footnotes/endnotes ───────────────────────────────────────
        footnotes, endnotes = self._extract_notes(doc)

        # ── Extract revision tracking ────────────────────────────────────────
        revisions = self._extract_revisions(doc)

        # ── Build text + lines for backward compatibility ────────────────────
        text_lines: list[str] = []
        for b in blocks:
            if b.text:
                text_lines.append(b.text)

        text = "\n".join(text_lines)

        # Build block_location_map + offset_map
        block_location_map: dict[int, SourceLocation] = {}
        from ..evidence.mapping import text_offset_mapper
        offset_map = text_offset_mapper.build_mapping(DocumentModel(
            filename=filename, content_type=content_type, text=text, lines=text_lines,
            headings=[], references=[], blocks=blocks,
        ))
        for i, b in enumerate(blocks):
            if b.location.line_start > 0:
                block_location_map[i] = b.location

        # Use TextParser for headings/references backward compat
        base = self.text_parser.parse(text, filename, content_type)

        # Build pages structure
        pages: list[list[int]] = []
        last_break = -1
        for pb in page_breaks:
            pages.append(list(range(last_break + 1, pb)))
            last_break = pb
        remaining = list(range(last_break + 1, len(blocks)))
        if remaining:
            pages.append(remaining)
        if not pages:
            pages = [list(range(len(blocks)))]

        return DocumentModel(
            filename=filename,
            content_type=content_type,
            text=text,
            lines=text_lines,
            headings=base.headings,
            references=base.references,
            blocks=blocks,
            tables=tables,
            figures=figures,
            metadata=metadata,
            block_location_map=block_location_map,
            footnotes=footnotes,
            endnotes=endnotes,
            revisions=revisions,
            page_breaks=page_breaks,
            pages=pages,
            offset_map=offset_map,
            hyperlinks=all_hyperlinks,
            citations=all_citations,
            sentences=all_sentences,
        )

    # ── Metadata ───────────────────────────────────────────────────────────────

    def _extract_metadata(self, doc: DocxDocument) -> DocumentMetadata:
        props = doc.core_properties
        # Page layout from sections
        page_width = page_height = 0.0
        margin_top = margin_bottom = margin_left = margin_right = 0.0
        default_font = ""
        default_font_size = 0.0
        line_spacing = 0.0

        if doc.sections:
            section = doc.sections[0]
            page_width = section.page_width or 0
            page_height = section.page_height or 0
            margin_top = section.top_margin or 0
            margin_bottom = section.bottom_margin or 0
            margin_left = section.left_margin or 0
            margin_right = section.right_margin or 0

        # Default font from styles
        try:
            default_style = doc.styles["Normal"]
            if default_style.font:
                default_font = default_style.font.name or ""
                if default_style.font.size:
                    default_font_size = default_style.font.size.pt or 0
            if default_style.paragraph_format:
                ls = default_style.paragraph_format.line_spacing
                if ls:
                    line_spacing = float(ls)
        except Exception:
            pass

        # Word count approximation
        word_count = 0
        char_count = 0
        for p in doc.paragraphs:
            words = p.text.split()
            word_count += len(words)
            char_count += len(p.text)

        return DocumentMetadata(
            title=str(props.title or ""),
            author=str(props.author or ""),
            subject=str(props.subject or ""),
            keywords=str(props.keywords or ""),
            creator=str(props.last_modified_by or props.author or ""),
            producer="Microsoft Word" if doc.element is not None else "",
            created_at=str(props.created or ""),
            modified_at=str(props.modified or ""),
            page_count=len(doc.sections[0].page) if doc.sections and hasattr(doc.sections[0], "page") else 1,
            word_count=word_count,
            character_count=char_count,
            page_width=page_width,
            page_height=page_height,
            margin_top=margin_top,
            margin_bottom=margin_bottom,
            margin_left=margin_left,
            margin_right=margin_right,
            default_font=default_font,
            default_font_size=default_font_size,
            line_spacing=line_spacing,
        )

    # ── Headers / Footers ─────────────────────────────────────────────────────

    def _extract_headers_footers(
        self, doc: DocxDocument
    ) -> tuple[list[DocumentBlock], list[DocumentBlock]]:
        headers: list[DocumentBlock] = []
        footers: list[DocumentBlock] = []

        for section in doc.sections:
            # Headers
            for header in (section.header, section.first_page_header, section.even_page_header):
                if header and not header.is_linked_to_previous:
                    for p in header.paragraphs:
                        if p.text.strip():
                            loc = SourceLocation(paragraph_id=self._get_paragraph_id(p))
                            headers.append(DocumentBlock(
                                type=BlockType.HEADER,
                                text=p.text.strip(),
                                location=loc,
                                style_name=p.style.name if p.style else "",
                                alignment=self._align_str(p.alignment),
                            ))
            # Footers
            for footer in (section.footer, section.first_page_footer, section.even_page_footer):
                if footer and not footer.is_linked_to_previous:
                    for p in footer.paragraphs:
                        if p.text.strip():
                            loc = SourceLocation(paragraph_id=self._get_paragraph_id(p))
                            footers.append(DocumentBlock(
                                type=BlockType.FOOTER,
                                text=p.text.strip(),
                                location=loc,
                                style_name=p.style.name if p.style else "",
                                alignment=self._align_str(p.alignment),
                            ))
        return headers, footers

    # ── Paragraph Parsing ─────────────────────────────────────────────────────

    def _parse_paragraph(
        self, p_element: Any, doc: DocxDocument, block_idx: int
    ) -> tuple[list[DocumentBlock], list[FigureData]]:
        blocks: list[DocumentBlock] = []
        figures: list[FigureData] = []

        # Check for page break
        for run in p_element.findall(".//w:br", self.NS):
            br_type = run.get(qn("w:type"), "")
            if br_type == "page":
                blocks.append(DocumentBlock(
                    type=BlockType.PAGE_BREAK, text="[PAGE BREAK]"
                ))

        # Get paragraph text with formatting
        p_text = ""
        p_runs = p_element.findall(".//w:r", self.NS)
        font_name = ""
        font_size = 0.0
        bold = False
        italic = False

        for r in p_runs:
            r_text = "".join(r.itertext())
            r_text = r_text.replace("\x00", "").strip("\r\n")

            # Get run formatting
            r_pr = r.find("w:rPr", self.NS)
            if r_pr is not None:
                r_font = r_pr.find("w:rFonts", self.NS)
                if r_font is not None:
                    font_name = r_font.get(qn("w:ascii"), "") or font_name
                r_sz = r_pr.find("w:sz", self.NS)
                if r_sz is not None:
                    try:
                        font_size = float(r_sz.get(qn("w:val"), 0)) / 2
                    except ValueError:
                        pass
                bold = r_pr.find("w:b", self.NS) is not None or bold
                italic = r_pr.find("w:i", self.NS) is not None or italic

            # Check for images in this run
            drawings = r.findall(".//w:drawing", self.NS)
            for drawing in drawings:
                fig = self._extract_figure(drawing)
                if fig:
                    figures.append(fig)

            p_text += r_text

        p_text = p_text.strip()
        if not p_text:
            return blocks, figures

        # Determine paragraph style
        p_pr = p_element.find("w:pPr", self.NS)
        style_name = ""
        heading_level = 0
        alignment = "left"
        indent = 0.0
        numbering = ""
        list_level = 0

        if p_pr is not None:
            # Style
            p_style = p_pr.find("w:pStyle", self.NS)
            if p_style is not None:
                style_name = p_style.get(qn("w:val"), "")

            # Heading detection
            if style_name.startswith("Heading"):
                try:
                    heading_level = int(style_name[-1])
                except ValueError:
                    pass
            # Also check for heading in paragraph properties
            outline_lvl = p_pr.find("w:outlineLvl", self.NS)
            if outline_lvl is not None:
                try:
                    heading_level = int(outline_lvl.get(qn("w:val"), 0)) + 1
                except ValueError:
                    pass

            # Alignment
            jc = p_pr.find("w:jc", self.NS)
            if jc is not None:
                alignment = jc.get(qn("w:val"), "left")

            # Indentation
            ind = p_pr.find("w:ind", self.NS)
            if ind is not None:
                left_val = ind.get(qn("w:left"), "0")
                try:
                    indent = float(left_val)
                except ValueError:
                    pass

            # Numbering
            num_pr = p_pr.find("w:numPr", self.NS)
            if num_pr is not None:
                num_id = num_pr.find("w:numId", self.NS)
                ilvl = num_pr.find("w:ilvl", self.NS)
                if num_id is not None:
                    numbering = num_id.get(qn("w:val"), "")
                if ilvl is not None:
                    try:
                        list_level = int(ilvl.get(qn("w:val"), 0))
                    except ValueError:
                        pass

        # Determine block type
        block_type = BlockType.PARAGRAPH
        if heading_level > 0:
            block_type = BlockType.HEADING
        elif numbering:
            block_type = BlockType.LIST_ITEM

        # Build location
        para_id = self._get_paragraph_id(p_element)
        loc = SourceLocation(
            paragraph_id=para_id,
            line_start=block_idx + 1,
            line_end=block_idx + 1,
        )

        blocks.append(DocumentBlock(
            type=block_type,
            text=p_text,
            level=heading_level,
            location=loc,
            font_name=font_name,
            font_size=font_size,
            bold=bold,
            italic=italic,
            alignment=alignment,
            style_name=style_name,
            list_type="numbered" if numbering else "bullet" if style_name == "List Bullet" else "",
            list_level=list_level,
            indent=indent,
        ))

        return blocks, figures

    # ── Figure Extraction ─────────────────────────────────────────────────────

    def _extract_figure(self, drawing: Any) -> FigureData | None:
        """Extract figure data from a drawing element."""
        try:
            # Try to get image from blip
            blip = drawing.find(".//a:blip", self.NS)
            if blip is None:
                return None

            # Get image dimensions
            ext = drawing.find(".//wp:extent", self.NS)
            img_width = 0.0
            img_height = 0.0
            if ext is not None:
                try:
                    img_width = float(ext.get("cx", 0)) / 914400  # EMU to inches
                    img_height = float(ext.get("cy", 0)) / 914400
                except ValueError:
                    pass

            # Get caption (usually next paragraph with "Figure" or "Table")
            caption = ""

            return FigureData(
                alt_text="",
                caption=caption,
                width=img_width,
                height=img_height,
            )
        except Exception:
            return None

    # ── Table Parsing ─────────────────────────────────────────────────────────

    def _parse_table(
        self, tbl_element: Any, doc: DocxDocument, block_idx: int
    ) -> tuple[TableData | None, list[DocumentBlock]]:
        try:
            # Get grid cols
            grid = tbl_element.find("w:tblGrid", self.NS)
            cols = 0
            col_widths: list[float] = []
            if grid is not None:
                for col in grid.findall("w:gridCol", self.NS):
                    cols += 1
                    try:
                        col_widths.append(float(col.get(qn("w:w"), 0)))
                    except ValueError:
                        col_widths.append(0)

            # Get table rows
            rows_elements = tbl_element.findall(".//w:tr", self.NS)
            rows_count = len(rows_elements)
            cells: list[TableCell] = []
            row_idx = 0

            for tr in rows_elements:
                tc_elements = tr.findall("w:tc", self.NS)
                col_idx = 0
                for tc in tc_elements:
                    # Merge info
                    tc_pr = tc.find("w:tcPr", self.NS)
                    colspan = 1
                    rowspan = 1
                    if tc_pr is not None:
                        grid_span = tc_pr.find("w:gridSpan", self.NS)
                        if grid_span is not None:
                            try:
                                colspan = int(grid_span.get(qn("w:val"), 1))
                            except ValueError:
                                pass
                        v_merge = tc_pr.find("w:vMerge", self.NS)
                        if v_merge is not None:
                            val = v_merge.get(qn("w:val"), "continue")
                            if val == "restart":
                                rowspan = 999  # Will be adjusted

                    # Cell text
                    cell_text = "".join(tc.itertext()).replace("\x00", "").strip()

                    cells.append(TableCell(
                        text=cell_text,
                        row=row_idx,
                        col=col_idx,
                        rowspan=rowspan,
                        colspan=colspan,
                        location=SourceLocation(
                            table_id=block_idx,
                            row_id=row_idx,
                            col_id=col_idx,
                        ),
                    ))
                    col_idx += colspan
                row_idx += 1

            # Get table caption (if any)
            caption = self._find_table_caption(tbl_element, doc)

            table_data = TableData(
                rows=rows_count,
                cols=cols or len(cells) // max(rows_count, 1),
                cells=cells,
                caption=caption,
                location=SourceLocation(table_id=block_idx),
            )

            # Create a block representation
            table_text = f"[Table: {rows_count} rows x {cols} cols]"
            if caption:
                table_text += f" — {caption}"

            blocks = [DocumentBlock(
                type=BlockType.TABLE,
                text=table_text,
                location=SourceLocation(table_id=block_idx),
                table=table_data,
            )]

            return table_data, blocks
        except Exception:
            return None, []

    def _find_table_caption(self, tbl_element: Any, doc: DocxDocument) -> str:
        """Look for a caption paragraph before or after the table."""
        body = doc.element.body
        tbl_index = list(body).index(tbl_element) if tbl_element in body else -1

        if tbl_index >= 0:
            # Check previous element
            if tbl_index > 0:
                prev = list(body)[tbl_index - 1]
                prev_text = "".join(prev.itertext()).strip()
                if prev_text and ("Table" in prev_text or "Bảng" in prev_text or "table" in prev_text.lower()):
                    return prev_text
            # Check next element
            if tbl_index + 1 < len(list(body)):
                next_el = list(body)[tbl_index + 1]
                next_text = "".join(next_el.itertext()).strip()
                if next_text and ("Table" in next_text or "Bảng" in next_text or "table" in next_text.lower()):
                    return next_text
        return ""

    # ── Footnotes / Endnotes ──────────────────────────────────────────────────

    def _extract_notes(self, doc: DocxDocument) -> tuple[list[str], list[str]]:
        footnotes: list[str] = []
        endnotes: list[str] = []

        try:
            # Access footnotes via relationship type
            for rel in doc.part.rels.values():
                reltype = rel.reltype.lower() if rel.reltype else ""
                if "footnote" in reltype and "footernote" not in reltype:
                    footnotes_xml = rel.target_part.blob
                    root = etree.fromstring(footnotes_xml)
                    for fn in root.findall(".//w:footnote", self.NS):
                        fn_id = fn.get(qn("w:id"), "")
                        fn_text = "".join(fn.itertext()).strip()
                        if fn_text and fn_id not in ("0", "-1"):
                            footnotes.append(fn_text)
                elif "endnote" in reltype:
                    endnotes_xml = rel.target_part.blob
                    root = etree.fromstring(endnotes_xml)
                    for en in root.findall(".//w:endnote", self.NS):
                        en_id = en.get(qn("w:id"), "")
                        en_text = "".join(en.itertext()).strip()
                        if en_text and en_id not in ("0", "-1"):
                            endnotes.append(en_text)
        except Exception:
            pass

        return footnotes, endnotes

    # ── Revision Tracking ─────────────────────────────────────────────────────

    def _extract_revisions(self, doc: DocxDocument) -> list[dict[str, Any]]:
        revisions: list[dict[str, Any]] = []

        try:
            body = doc.element.body
            ns = self.NS

            # Find all tracked changes (insertions / deletions)
            for ins in body.findall(".//w:ins", ns):
                author = ins.get(qn("w:author"), "")
                date = ins.get(qn("w:date"), "")
                text = "".join(ins.itertext()).strip()
                if text:
                    revisions.append({
                        "type": "insertion",
                        "author": author,
                        "date": date,
                        "text": text[:200],
                    })

            for dele in body.findall(".//w:del", ns):
                author = dele.get(qn("w:author"), "")
                date = dele.get(qn("w:date"), "")
                text = "".join(dele.itertext()).strip()
                if text:
                    revisions.append({
                        "type": "deletion",
                        "author": author,
                        "date": date,
                        "text": text[:200],
                    })
        except Exception:
            pass

        return revisions

    # ── SDT (Content Controls) ────────────────────────────────────────────────

    def _parse_sdt(
        self, sdt_element: Any, doc: DocxDocument, block_idx: int
    ) -> tuple[list[DocumentBlock], list[FigureData], list[TableData]]:
        blocks: list[DocumentBlock] = []
        figures: list[FigureData] = []
        tables: list[TableData] = []

        # Process paragraphs inside SDT
        for p in sdt_element.findall(".//w:p", self.NS):
            p_blocks, p_figures = self._parse_paragraph(p, doc, block_idx + len(blocks))
            blocks.extend(p_blocks)
            figures.extend(p_figures)

        # Process tables inside SDT
        for t in sdt_element.findall(".//w:tbl", self.NS):
            table_data, table_blocks = self._parse_table(t, doc, block_idx + len(blocks))
            if table_data:
                tables.append(table_data)
            blocks.extend(table_blocks)

        return blocks, figures, tables

    # ── Hyperlink & Citation Extraction ────────────────────────────────────────

    def _extract_hyperlinks_and_citations(
        self, p_element: Any, block: DocumentBlock, doc: Any = None
    ) -> tuple[list[Hyperlink], list[Citation]]:
        """Extract hyperlinks and citations from a paragraph."""
        hyperlinks: list[Hyperlink] = []
        citations: list[Citation] = []

        # Extract hyperlinks from w:hyperlink elements
        for hl in p_element.findall(".//w:hyperlink", self.NS):
            hl_id = hl.get("r:id", "") or hl.get(qn("w:anchor"), "")
            hl_text = "".join(hl.itertext()).strip()
            if hl_text and hl_id:
                # Resolve URL via document relationships
                resolved_url = hl_id
                if doc and hasattr(doc, 'part') and hasattr(doc.part, 'rels'):
                    rel = doc.part.rels.get(hl_id)
                    if rel:
                        resolved_url = rel.target_ref or hl_id
                hyperlinks.append(Hyperlink(
                    text=hl_text,
                    url=resolved_url,
                    location=block.location,
                ))

        # Extract citations via field codes (w:instrText for CITATION, REF fields)
        for field in p_element.findall(".//w:instrText", self.NS):
            field_text = field.text or ""
            # CITATION field
            if "CITATION" in field_text:
                citations.append(Citation(
                    raw_text=field_text,
                    citation_type="inline",
                    location=block.location,
                ))
            # REF field
            elif " REF " in field_text:
                citations.append(Citation(
                    raw_text=field_text,
                    citation_type="cross_ref",
                    location=block.location,
                ))
            # BIBLIOGRAPHY field
            elif "BIBLIOGRAPHY" in field_text or "BIBLIO" in field_text:
                citations.append(Citation(
                    raw_text=field_text,
                    citation_type="biblio",
                    location=block.location,
                ))

        # Extract bracketed citations like [1], [Author, Year]
        import re
        bracket_refs = re.findall(r"\[(\d+(?:,\s*\d+)*)\]", block.text)
        for ref in bracket_refs:
            citations.append(Citation(
                raw_text=f"[{ref}]",
                citation_type="bracket",
                location=block.location,
            ))

        # APA-style (Author, Year)
        apa_refs = re.findall(r"\(([A-Z][a-z]+(?:\s+et\s+al\.)?,\s*\d{4})\)", block.text)
        for ref in apa_refs:
            citations.append(Citation(
                raw_text=f"({ref})",
                citation_type="apa",
                location=block.location,
            ))

        return hyperlinks, citations

    # ── Helpers ───────────────────────────────────────────────────────────────

    @staticmethod
    def _get_paragraph_id(p_element: Any) -> int:
        """Get paragraph ID from the <w:p> element's w:paraId attribute."""
        try:
            # w:paraId is an attribute ON the paragraph element, not inside w:pPr
            para_id = p_element.get(qn("w:paraId"), None) if hasattr(p_element, "get") else None
            if para_id is not None:
                return int(para_id, 16) & 0xFFFF
        except Exception:
            pass
        return 0

    @staticmethod
    def _align_str(alignment: Any) -> str:
        """Convert WD_ALIGN_PARAGRAPH to string."""
        mapping = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        }
        return mapping.get(alignment, "left")
