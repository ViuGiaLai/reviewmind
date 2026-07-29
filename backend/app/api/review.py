from __future__ import annotations

import csv
import io
import json
import threading
from dataclasses import asdict, replace
from datetime import datetime, timezone
from typing import Any

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import HTMLResponse, Response
from pydantic import BaseModel, Field

from app.api.auth import get_current_user_id, require_resource_owner
from app.config import settings
from app.database import create_database
from app.operations import metrics
from app.review import ReviewEngine, ReviewRequest
from app.review.models import ReviewResult
from app.review.parser import FileParser
from app.storage import create_storage
from app.security import audit_log

router = APIRouter(prefix="/api", tags=["review"])

engine = ReviewEngine()
file_parser = FileParser()
database = create_database()
storage = create_storage()
_review_slots = threading.BoundedSemaphore(settings.performance.max_concurrent_reviews)

# ─── Schemas ───────────────────────────────────────────────────────────────────

class ReviewBody(BaseModel):
    text: str = Field(default="", min_length=0, max_length=2_000_000)
    filename: str = "document.md"
    content_type: str = "text/markdown"
    profile_id: str = "academic"
    pack_ids: list[str] = []
    enabled_categories: list[str] | None = None
    review_mode: str = Field(default="rule_ai", pattern=r"^(rule_only|rule_ai|full)$")
    report_language: str = Field(default="en", pattern=r"^(en|vi)$")
    document_id: str | None = None
    template_id: str | None = None


class IssueUpdateBody(BaseModel):
    status: str = Field(pattern=r"^(open|resolved|ignored)$")


class BulkIssueUpdateBody(BaseModel):
    status: str = Field(pattern=r"^(open|resolved|ignored)$")
    category: str | None = None


class ExportQuery(BaseModel):
    format: str = Field(default="json", pattern=r"^(json|md|markdown|html|pdf|docx|csv)$")


# ─── Helper ────────────────────────────────────────────────────────────────────

def _verify_session_owner(session_id: str, user_id: str | None) -> dict[str, Any]:
    """Return a session only when it belongs to the active account."""
    return require_resource_owner(database.get_session(session_id), user_id, "Session not found.")


def _run_review(
    text: str,
    filename: str,
    content_type: str,
    profile_id: str,
    pack_ids: list[str],
    enabled_categories: list[str] | None,
    review_mode: str = "rule_ai",
    report_language: str = "en",
    source_content: bytes | None = None,
    reference_template: dict[str, Any] | None = None,
    reference_template_id: str = "",
    profile_overrides: dict[str, Any] | None = None,
    scoring_mode: str = "standard",
    auto_fix_enabled: bool = True,
) -> tuple[dict[str, Any], ReviewResult]:
    """Run a review and return both the dict form and the domain object."""
    request = ReviewRequest(
        text=text,
        filename=filename,
        content_type=content_type,
        profile_id=profile_id,
        pack_ids=pack_ids,
        enabled_categories=enabled_categories,
        review_mode=review_mode,
        report_language=report_language,
        source_content=source_content,
        reference_template=reference_template,
        reference_template_id=reference_template_id,
        profile_overrides=profile_overrides or {},
        scoring_mode=scoring_mode,
    )
    acquired = _review_slots.acquire(
        timeout=settings.performance.review_queue_timeout_seconds
    )
    if not acquired:
        metrics.increment("reviewmind_review_rejected_total", reason="capacity")
        raise HTTPException(
            status_code=503,
            detail="Review capacity is temporarily full. Retry shortly.",
            headers={"Retry-After": "2"},
        )
    metrics.increment("reviewmind_reviews_inflight", value=1)
    try:
        result = engine.review(request)
        if not auto_fix_enabled:
            result = replace(
                result,
                issues=[replace(issue, autofix_allowed=False) for issue in result.issues],
            )
        result_dict = asdict(result)
        return result_dict, result
    finally:
        metrics.increment("reviewmind_reviews_inflight", value=-1)
        _review_slots.release()


# ═══════════════════════════════════════════════════════════════════════════════
# 1. TEXT REVIEW (existing)
# ═══════════════════════════════════════════════════════════════════════════════

@router.get("/reviews")
def list_reviews(
    limit: int = Query(30, ge=1, le=100),
    offset: int = Query(0, ge=0),
    user_id: str | None = Depends(get_current_user_id),
) -> dict[str, Any]:
    """List review sessions for the current user."""
    sessions, total = database.list_sessions(limit=limit, offset=offset, user_id=user_id)
    return {"reviews": sessions, "total": total}


@router.post("/reviews")
def review_text(
    body: ReviewBody,
    user_id: str | None = Depends(get_current_user_id),
) -> dict[str, Any]:
    """Review a document — either from raw text or from a previously uploaded document (by document_id)."""
    try:
        doc_id = body.document_id
        filename = body.filename
        content_type = body.content_type
        source_content: bytes | None = None
        template_record: dict[str, Any] | None = None

        execution_profile_id = body.profile_id
        effective_pack_ids = body.pack_ids
        effective_categories = body.enabled_categories
        effective_review_mode = body.review_mode
        effective_language = body.report_language
        effective_template_id = body.template_id
        profile_overrides: dict[str, Any] = {}
        scoring_mode = "standard"
        auto_fix_enabled = True

        try:
            engine.profiles.load(body.profile_id)
        except ValueError:
            custom_profile = require_resource_owner(
                database.get_evaluation_profile(body.profile_id), user_id,
                "Evaluation profile not found.",
            )
            execution_profile_id = custom_profile["base_profile_id"]
            effective_pack_ids = body.pack_ids or custom_profile["knowledge_pack_ids"]
            effective_categories = body.enabled_categories or custom_profile["enabled_categories"]
            effective_template_id = body.template_id or custom_profile.get("reference_template_id")
            effective_language = custom_profile["language"]
            scoring_mode = custom_profile["review_mode"]
            auto_fix_enabled = custom_profile["auto_fix_enabled"]
            if not custom_profile["ai_review_enabled"]:
                effective_review_mode = "rule_only"
            base_profile = engine.profiles.load(execution_profile_id)
            categories = effective_categories or base_profile.categories
            weights = (
                {category: 1 for category in categories}
                if custom_profile["scoring_profile"] == "equal"
                else {category: base_profile.weights.get(category, 1) for category in categories}
            )
            profile_overrides = {
                "id": custom_profile["id"],
                "name": custom_profile["name"],
                "categories": categories,
                "weights": weights,
            }

        if effective_template_id:
            template_record = require_resource_owner(
                database.get_reference_template(effective_template_id), user_id,
                "Reference template not found.",
            )

        # ── Resolve document content ──
        if doc_id:
            # Fetch from storage + parse
            doc_record = require_resource_owner(
                database.get_document(doc_id), user_id, "Document not found."
            )
            content = storage.read(doc_record["storage_path"])
            source_content = content
            filename = doc_record["original_name"]
            content_type = doc_record.get("content_type", "application/octet-stream")
            document = file_parser.parse(content, filename, content_type)
            text = document.text
        else:
            # Use provided text directly
            if not body.text:
                raise HTTPException(status_code=400, detail="Provide either 'text' or 'document_id'.")
            text = body.text

        # ── Run review ──
        result_dict, result_obj = _run_review(
            text=text,
            filename=filename,
            content_type=content_type,
            profile_id=execution_profile_id,
            pack_ids=effective_pack_ids,
            enabled_categories=effective_categories,
            review_mode=effective_review_mode,
            report_language=effective_language,
            source_content=source_content,
            reference_template=template_record.get("analysis") if template_record else None,
            reference_template_id=effective_template_id or "",
            profile_overrides=profile_overrides,
            scoring_mode=scoring_mode,
            auto_fix_enabled=auto_fix_enabled,
        )

        # Persist raw-text reviews as documents too. Auto Fix retrieves its
        # immutable source through session.document_id, so leaving this null
        # made the primary paste-text flow impossible to fix or version.
        if not doc_id:
            encoded = text.encode("utf-8")
            storage_path, _ = storage.save(encoded, filename)
            doc_id = database.save_document(
                original_name=filename,
                content_type=content_type,
                size=len(encoded),
                storage_path=storage_path,
                user_id=user_id,
            )

        session_id = database.save_session(
            filename=filename,
            profile_id=body.profile_id,
            pack_ids=effective_pack_ids,
            categories=list(result_dict["category_scores"].keys()),
            result=result_obj,
            document_id=doc_id,
            user_id=user_id,
            reference_template_id=effective_template_id,
        )
        audit_log.record(
            actor_id=user_id or "anonymous", action="review.completed",
            resource_type="review_session", resource_id=session_id,
            metadata={"profile_id": body.profile_id, "score": result_obj.score},
        )
        result_dict["session_id"] = session_id
        result_dict["document_id"] = doc_id
        result_dict["filename"] = filename
        result_dict["template_id"] = effective_template_id
        result_dict["template_name"] = template_record.get("original_name") if template_record else None

        # Replace issues with DB-stored issues so they have real UUIDs,
        # not the engine's rule_id-style IDs.
        db_issues, _ = database.list_issues(session_id=session_id, limit=500)
        result_dict["issues"] = db_issues

        return result_dict
    except HTTPException:
        raise
    except ValueError as error:
        raise HTTPException(status_code=422, detail=str(error)) from error


# ═══════════════════════════════════════════════════════════════════════════════
# 3. HISTORY
# ═══════════════════════════════════════════════════════════════════════════════

@router.get("/history")
def list_history(
    profile_id: str | None = Query(None),
    document_id: str | None = Query(None),
    status: str | None = Query(None),
    score_min: int | None = Query(None, ge=0, le=100),
    score_max: int | None = Query(None, ge=0, le=100),
    search: str | None = Query(None),
    limit: int = Query(30, ge=1, le=100),
    offset: int = Query(0, ge=0),
    user_id: str | None = Depends(get_current_user_id),
) -> dict[str, Any]:
    """List review sessions with filtering and pagination (scoped to current user)."""
    sessions, total = database.list_sessions(
        profile_id=profile_id,
        document_id=document_id,
        status=status,
        score_min=score_min,
        score_max=score_max,
        search=search,
        limit=limit,
        offset=offset,
        user_id=user_id,
    )
    return {
        "items": sessions,
        "total": total,
        "limit": limit,
        "offset": offset,
    }


@router.get("/history/{session_id}")
def get_session_detail(
    session_id: str,
    user_id: str | None = Depends(get_current_user_id),
) -> dict[str, Any]:
    """Get full detail of a review session including issues."""
    session = _verify_session_owner(session_id, user_id)


    issues, issue_count = database.list_issues(session_id=session_id, limit=500)
    session["issues"] = issues
    session["issue_count"] = issue_count

    # Get document info if available
    if session.get("document_id"):
        doc = require_resource_owner(database.get_document(session["document_id"]), user_id, "Document not found.")
        session["document"] = doc

    return session


@router.delete("/history/{session_id}")
def delete_session(
    session_id: str,
    user_id: str | None = Depends(get_current_user_id),
) -> dict[str, str]:
    """Delete a review session and its issues."""
    session = _verify_session_owner(session_id, user_id)

    if not database.delete_session(session_id):
        raise HTTPException(status_code=404, detail="Session not found.")
    return {"status": "deleted"}


# ═══════════════════════════════════════════════════════════════════════════════
# 4. ISSUES
# ═══════════════════════════════════════════════════════════════════════════════

@router.get("/sessions/{session_id}/issues")
def list_session_issues(
    session_id: str,
    severity: str | None = Query(None, pattern=r"^(high|medium|low)$"),
    category: str | None = Query(None),
    status: str | None = Query(None, pattern=r"^(open|resolved|ignored)$"),
    rule_id: str | None = Query(None),
    search: str | None = Query(None),
    limit: int = Query(100, ge=1, le=500),
    offset: int = Query(0, ge=0),
    user_id: str | None = Depends(get_current_user_id),
) -> dict[str, Any]:
    """List issues for a session with filtering and pagination."""
    # Verify session exists + ownership
    session = _verify_session_owner(session_id, user_id)

    issues, total = database.list_issues(
        session_id=session_id,
        severity=severity,
        category=category,
        status=status,
        rule_id=rule_id,
        search=search,
        limit=limit,
        offset=offset,
    )
    return {
        "items": issues,
        "total": total,
        "limit": limit,
        "offset": offset,
        "session_score": session.get("score"),
    }


@router.get("/sessions/{session_id}/insights")
def get_session_insights(
    session_id: str,
    user_id: str | None = Depends(get_current_user_id),
) -> dict[str, Any]:
    """Get rich insights and recommendations for a session."""
    session = _verify_session_owner(session_id, user_id)
    issues, _ = database.list_issues(session_id=session_id, limit=500)

    from app.review.report.insights import InsightsEngine
    from app.review.profiles import ProfileLoader
    from dataclasses import dataclass

    # We need to construct mock Issue objects for the InsightsEngine
    @dataclass
    class MockIssue:
        category: str
        severity: Any
        rule_id: str
        message: str
        recommendation: str
        autofix_allowed: bool
        source: str

    @dataclass
    class MockSeverity:
        value: str

    mock_issues = []
    for issue in issues:
        mock_issues.append(MockIssue(
            category=issue.get("category", "other"),
            severity=MockSeverity(value=issue.get("severity", "medium")),
            rule_id=issue.get("rule_id", ""),
            message=issue.get("message", ""),
            recommendation=issue.get("recommendation", ""),
            autofix_allowed=bool(issue.get("autofix_allowed", False)),
            source=issue.get("source", "rule"),
        ))

    # Get profile
    try:
        from pathlib import Path
        config_dir = Path(__file__).resolve().parents[3] / "config"
        loader = ProfileLoader(config_dir)
        profile = loader.load(session.get("profile_id", "academic"))
    except Exception:
        profile = None

    engine = InsightsEngine()
    report = engine.generate_report(
        issues=mock_issues,
        score=session.get("score", 0),
        category_scores=session.get("category_scores", {}),
        profile=profile,
    )

    from dataclasses import asdict, replace
    return asdict(report)


@router.get("/sessions/{session_id}/issues/stats")
def get_issue_stats(
    session_id: str,
    user_id: str | None = Depends(get_current_user_id),
) -> dict[str, Any]:
    """Get aggregated statistics for issues in a session."""
    session = _verify_session_owner(session_id, user_id)

    issues, total = database.list_issues(session_id=session_id, limit=500)

    severity_counts: dict[str, int] = {"high": 0, "medium": 0, "low": 0}
    category_counts: dict[str, int] = {}
    status_counts: dict[str, int] = {"open": 0, "resolved": 0, "ignored": 0}

    for issue in issues:
        sev = issue.get("severity", "low")
        severity_counts[sev] = severity_counts.get(sev, 0) + 1
        cat = issue.get("category", "unknown")
        category_counts[cat] = category_counts.get(cat, 0) + 1
        sts = issue.get("status", "open")
        status_counts[sts] = status_counts.get(sts, 0) + 1

    return {
        "total": total,
        "by_severity": severity_counts,
        "by_category": category_counts,
        "by_status": status_counts,
    }


@router.get("/sessions/{session_id}/issues/{issue_id}")
def get_issue_detail(
    session_id: str,
    issue_id: str,
    user_id: str | None = Depends(get_current_user_id),
) -> dict[str, Any]:
    """Get full detail of a single issue, including history across scans."""
    _verify_session_owner(session_id, user_id)
    issue = database.get_issue(issue_id)
    if not issue:
        raise HTTPException(status_code=404, detail="Issue not found.")
    if issue.get("session_id") != session_id:
        raise HTTPException(status_code=404, detail="Issue not found in this session.")

    # Get history of same issue pattern across previous scans
    history = database.get_issue_history(issue["issue_id"], session_id, user_id=user_id)
    issue["scan_history"] = history

    return issue


@router.patch("/sessions/{session_id}/issues/{issue_id}")
def update_issue_status(
    session_id: str,
    issue_id: str,
    body: IssueUpdateBody,
    user_id: str | None = Depends(get_current_user_id),
) -> dict[str, Any]:
    """Update the status of an issue (open/resolved/ignored)."""
    _verify_session_owner(session_id, user_id)
    issue = database.get_issue(issue_id)
    if not issue:
        raise HTTPException(status_code=404, detail="Issue not found.")
    if issue.get("session_id") != session_id:
        raise HTTPException(status_code=404, detail="Issue not found in this session.")

    database.update_issue_status(issue_id, body.status)
    return {"id": issue_id, "status": body.status}


@router.post("/sessions/{session_id}/issues/bulk")
def bulk_update_issues(
    session_id: str,
    body: BulkIssueUpdateBody,
    user_id: str | None = Depends(get_current_user_id),
) -> dict[str, Any]:
    """Bulk update issue status for a session."""
    _verify_session_owner(session_id, user_id)
    count = database.bulk_update_issue_status(session_id, body.status, body.category)
    return {"updated": count, "status": body.status, "category": body.category}


# ═══════════════════════════════════════════════════════════════════════════════
# 5. EXPORT
# ═══════════════════════════════════════════════════════════════════════════════

@router.get("/sessions/{session_id}/export")
def export_session(
    session_id: str,
    format: str = Query("json", pattern=r"^(json|md|markdown|html|pdf|docx|csv)$"),
    user_id: str | None = Depends(get_current_user_id),
) -> dict[str, Any]:
    """Export a review session in the requested format."""
    session = _verify_session_owner(session_id, user_id)

    issues, _ = database.list_issues(session_id=session_id, limit=500)

    if format == "json":
        data = {
            "session": {k: v for k, v in session.items() if k != "issues"},
            "issues": issues,
            "exported_at": datetime.now(timezone.utc).isoformat(),
        }
        return Response(
            content=json.dumps(data, indent=2, ensure_ascii=False),
            media_type="application/json",
            headers={"Content-Disposition": f'attachment; filename="review_{session_id[:8]}.json"'},
        )

    elif format in ("md", "markdown"):
        content = _export_markdown(session, issues)
        return Response(
            content=content,
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="review_{session_id[:8]}.md"'},
        )

    elif format == "html":
        content = _export_html(session, issues)
        return HTMLResponse(
            content=content,
            headers={"Content-Disposition": f'attachment; filename="review_{session_id[:8]}.html"'},
        )

    elif format == "csv":
        content = _export_csv(session, issues)
        return Response(
            content=content,
            media_type="text/csv",
            headers={"Content-Disposition": f'attachment; filename="review_{session_id[:8]}.csv"'},
        )

    elif format == "pdf":
        return _export_pdf(session, issues, session_id)

    elif format == "docx":
        return _export_docx(session, issues, session_id)

    raise HTTPException(status_code=400, detail=f"Unsupported export format: {format}")


def _export_markdown(session: dict[str, Any], issues: list[dict[str, Any]]) -> str:
    lines = [
        f"# ReviewMind Report",
        "",
        f"**Session:** {session['id'][:8]}",
        f"**Profile:** {session['profile_id']}",
        f"**Document:** {session['filename']}",
        f"**Date:** {session.get('created_at', 'N/A')}",
        f"**Overall Score:** {session['score']}/100",
        "",
        "## Category Scores",
    ]
    cat_scores = session.get("category_scores", {})
    for cat, score in cat_scores.items():
        bars = "█" * (score // 10) + "░" * (10 - score // 10)
        lines.append(f"- **{cat}:** {score}/100  `{bars}`")

    lines.extend(["", "## Summary", "", session.get("summary", ""), "", "## Issues"])

    if not issues:
        lines.append("No issues found.")
    else:
        for i, issue in enumerate(issues, 1):
            severity_badge = {
                "high": "🔴 HIGH",
                "medium": "🟡 MEDIUM",
                "low": "🟢 LOW",
            }.get(issue.get("severity", ""), issue.get("severity", ""))
            lines.extend([
                f"### {i}. {severity_badge} — {issue['message']}",
                f"- **Category:** {issue.get('category', 'N/A')}",
                f"- **Rule:** `{issue.get('rule_id', 'N/A')}`",
                f"- **Confidence:** {issue.get('confidence', 0)}%",
                f"- **Status:** {issue.get('status', 'open')}",
                f"- **Evidence:** \"{issue.get('evidence_excerpt', '')}\"",
                f"- **Location:** {issue.get('evidence_location', 'N/A')}",
                f"- **Recommendation:** {issue.get('recommendation', 'N/A')}",
                "",
            ])

    lines.extend([
        "---",
        f"*Generated by ReviewMind on {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}*",
    ])
    return "\n".join(lines)


def _export_html(session: dict[str, Any], issues: list[dict[str, Any]]) -> str:
    bars = ""
    for cat, score in session.get("category_scores", {}).items():
        pct = min(score, 100)
        bars += f"<div class='cat-bar'><span>{cat}</span><div class='bar'><div class='fill' style='width:{pct}%'></div></div><strong>{score}</strong></div>"

    issue_rows = ""
    if not issues:
        issue_rows = "<p class='muted'>No issues found.</p>"
    else:
        for issue in issues:
            sev = issue.get("severity", "low")
            icon = {"high": "🔴", "medium": "🟡", "low": "🟢"}.get(sev, "")
            issue_rows += f"""
            <div class="issue {sev}">
                <div class="issue-header">
                    <span class="badge {sev}">{icon} {sev.upper()}</span>
                    <span class="category">{issue.get('category', '')}</span>
                    <span class="confidence">{issue.get('confidence', 0)}% confidence</span>
                    <span class="status">{issue.get('status', 'open')}</span>
                </div>
                <h4>{issue['message']}</h4>
                <blockquote>"{issue.get('evidence_excerpt', '')}"</blockquote>
                <p class="rec"><strong>Recommendation:</strong> {issue.get('recommendation', '')}</p>
                <p class="meta">Rule: <code>{issue.get('rule_id', '')}</code> &middot; Location: {issue.get('evidence_location', '')}</p>
            </div>"""

    score = session.get("score", 0)
    score_color = "#22c55e" if score >= 80 else "#eab308" if score >= 50 else "#ef4444"

    return f"""<!DOCTYPE html>
<html lang="en">
<head><meta charset="utf-8"><title>ReviewMind Report — {session['id'][:8]}</title>
<style>
* {{ box-sizing: border-box; margin: 0; padding: 0; }}
body {{ font-family: Inter, system-ui, sans-serif; background: #f5f7f8; color: #192324; padding: 40px 24px; }}
.container {{ max-width: 900px; margin: auto; }}
.header {{ text-align: center; margin-bottom: 32px; }}
.score-ring {{ width: 120px; height: 120px; border-radius: 50%; background: conic-gradient({score_color} {score * 3.6}deg, #e5e7eb 0deg); display: flex; align-items: center; justify-content: center; margin: 16px auto; }}
.score-ring span {{ background: white; width: 90px; height: 90px; border-radius: 50%; display: flex; align-items: center; justify-content: center; font-size: 2rem; font-weight: 800; color: {score_color}; }}
.meta {{ color: #536069; font-size: .9rem; }}
h2 {{ margin: 24px 0 12px; }}
.cat-bar {{ display: flex; align-items: center; gap: 12px; margin: 8px 0; }}
.cat-bar span {{ width: 100px; text-transform: capitalize; }}
.bar {{ flex: 1; height: 12px; background: #e5e7eb; border-radius: 6px; }}
.fill {{ height: 100%; border-radius: 6px; background: #18716a; }}
.issue {{ border: 1px solid #e1e7e8; border-radius: 12px; padding: 20px; margin: 16px 0; background: white; }}
.issue.high {{ border-left: 5px solid #ef4444; }}
.issue.medium {{ border-left: 5px solid #eab308; }}
.issue.low {{ border-left: 5px solid #3b82f6; }}
.issue-header {{ display: flex; gap: 12px; align-items: center; flex-wrap: wrap; margin-bottom: 8px; font-size: .85rem; }}
.badge {{ font-weight: 700; padding: 2px 8px; border-radius: 4px; }}
.badge.high {{ background: #fef2f2; color: #dc2626; }}
.badge.medium {{ background: #fefce8; color: #ca8a04; }}
.badge.low {{ background: #eff6ff; color: #2563eb; }}
.category {{ text-transform: capitalize; color: #536069; }}
.confidence {{ color: #718087; }}
.status {{ text-transform: capitalize; background: #f0f5f5; padding: 2px 8px; border-radius: 4px; }}
.issue h4 {{ margin: 4px 0 8px; font-size: 1.05rem; }}
blockquote {{ border-left: 3px solid #cbd5d8; padding-left: 12px; color: #55636b; margin: 8px 0; font-style: italic; }}
.rec {{ margin: 8px 0; }}
.meta {{ font-size: .85rem; color: #718087; }}
.muted {{ color: #718087; text-align: center; padding: 24px; }}
hr {{ border: 0; border-top: 1px solid #e1e7e8; margin: 32px 0; }}
</style></head>
<body><div class="container">
<div class="header">
    <h1>ReviewMind</h1>
    <p class="meta">Document Review Report</p>
    <p class="meta">{session.get('filename', 'N/A')} &middot; {session.get('profile_id', 'N/A')} profile</p>
    <div class="score-ring"><span>{score}</span></div>
    <p class="meta">{session.get('summary', '')}</p>
</div>
<h2>Category Scores</h2>
{bars}
<h2>Issues ({len(issues)})</h2>
{issue_rows}
<hr>
<p class="meta">Generated by ReviewMind on {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}</p>
</div></body></html>"""


def _export_csv(session: dict[str, Any], issues: list[dict[str, Any]]) -> str:
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Session ID", "Profile", "Document", "Date", "Score", "Summary"])
    writer.writerow([
        session["id"][:8],
        session["profile_id"],
        session["filename"],
        session.get("created_at", ""),
        session["score"],
        session.get("summary", ""),
    ])
    writer.writerow([])
    writer.writerow(["Issue ID", "Severity", "Category", "Rule", "Message", "Confidence",
                     "Status", "Location", "Recommendation", "Evidence"])

    for issue in issues:
        writer.writerow([
            issue.get("issue_id", ""),
            issue.get("severity", ""),
            issue.get("category", ""),
            issue.get("rule_id", ""),
            issue.get("message", ""),
            issue.get("confidence", 0),
            issue.get("status", ""),
            issue.get("evidence_location", ""),
            issue.get("recommendation", ""),
            issue.get("evidence_excerpt", ""),
        ])

    return output.getvalue()


def _export_pdf(session: dict[str, Any], issues: list[dict[str, Any]], session_id: str) -> Response:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable,
        )
        from reportlab.lib.enums import TA_CENTER, TA_LEFT

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                leftMargin=20*mm, rightMargin=20*mm,
                                topMargin=20*mm, bottomMargin=20*mm)

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("Title2", parent=styles["Title"], fontSize=24, spaceAfter=6)
        subtitle_style = ParagraphStyle("Sub", parent=styles["Normal"], fontSize=10, textColor=colors.grey, alignment=TA_CENTER)
        heading2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=14, spaceBefore=16, spaceAfter=8)
        heading3 = ParagraphStyle("H3", parent=styles["Heading3"], fontSize=11, spaceBefore=10, spaceAfter=4)
        normal = ParagraphStyle("Normal2", parent=styles["Normal"], fontSize=9, leading=13, spaceAfter=4)
        bold = ParagraphStyle("Bold", parent=normal, fontName="Helvetica-Bold")
        code_style = ParagraphStyle("Code", parent=normal, fontName="Courier", fontSize=8)

        elements = []

        # Title
        elements.append(Paragraph("ReviewMind Report", title_style))
        elements.append(Paragraph(
            f"Session: {session['id'][:8]} &nbsp;|&nbsp; Profile: {session['profile_id']} &nbsp;|&nbsp; Document: {session.get('filename', 'N/A')}",
            subtitle_style,
        ))
        elements.append(Spacer(1, 8*mm))

        # Score
        score = session.get("score", 0)
        score_color = colors.green if score >= 80 else (colors.orange if score >= 50 else colors.red)
        score_style = ParagraphStyle("Score", parent=styles["Normal"], fontSize=36, textColor=score_color, alignment=TA_CENTER)
        elements.append(Paragraph(f"{score}/100", score_style))
        elements.append(Paragraph("Overall Quality Score", subtitle_style))
        elements.append(Spacer(1, 6*mm))

        # Category scores table
        cat_scores = session.get("category_scores", {})
        if cat_scores:
            elements.append(Paragraph("Category Scores", heading2))
            cat_data = [["Category", "Score"]] + [[cat.capitalize(), str(sc)] for cat, sc in cat_scores.items()]
            cat_table = Table(cat_data, colWidths=[120, 60])
            cat_table.setStyle(TableStyle([
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#18716a")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("ALIGN", (1, 0), (-1, -1), "CENTER"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f5f5")]),
            ]))
            elements.append(cat_table)
            elements.append(Spacer(1, 4*mm))

        # Summary
        elements.append(Paragraph("Summary", heading2))
        elements.append(Paragraph(session.get("summary", ""), normal))
        elements.append(Spacer(1, 4*mm))

        # Issues
        elements.append(Paragraph(f"Issues ({len(issues)})", heading2))
        if not issues:
            elements.append(Paragraph("No issues found.", normal))
        else:
            for issue in issues:
                sev = issue.get("severity", "low")
                sev_color = {"high": "#ef4444", "medium": "#eab308", "low": "#3b82f6"}.get(sev, "#666")
                elements.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e1e7e8")))
                elements.append(Paragraph(
                    f'<font color="{sev_color}">{"🔴" if sev=="high" else "🟡" if sev=="medium" else "🟢"} {sev.upper()}</font> &nbsp;—&nbsp; {issue["message"]}',
                    bold,
                ))
                elements.append(Paragraph(
                    f"Category: <b>{issue.get('category', '')}</b> &nbsp;|&nbsp; "
                    f"Rule: <font face='Courier'>{issue.get('rule_id', '')}</font> &nbsp;|&nbsp; "
                    f"Confidence: {issue.get('confidence', 0)}% &nbsp;|&nbsp; "
                    f"Status: {issue.get('status', 'open')}",
                    code_style,
                ))
                if issue.get("evidence_excerpt"):
                    elements.append(Paragraph(
                        f"<i>Evidence:</i> &ldquo;{issue['evidence_excerpt'][:200]}&rdquo;",
                        normal,
                    ))
                elements.append(Paragraph(
                    f"<b>Recommendation:</b> {issue.get('recommendation', '')}",
                    normal,
                ))
                elements.append(Spacer(1, 2*mm))

        # Footer
        elements.append(Spacer(1, 8*mm))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cbd5d8")))
        elements.append(Paragraph(
            f"Generated by ReviewMind on {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
            ParagraphStyle("Footer", parent=normal, fontSize=8, textColor=colors.grey, alignment=TA_CENTER),
        ))

        doc.build(elements)
        pdf_content = buffer.getvalue()
        buffer.close()

        return Response(
            content=pdf_content,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="review_{session_id[:8]}.pdf"'},
        )
    except ImportError:
        raise HTTPException(status_code=501, detail="PDF export requires reportlab library.")
    except Exception as error:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {error}")


def _export_docx(session: dict[str, Any], issues: list[dict[str, Any]], session_id: str) -> Response:
    try:
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        doc = DocxDocument()

        # Title
        title = doc.add_heading("ReviewMind Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Meta
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta.add_run(
            f"Session: {session['id'][:8]}  |  Profile: {session['profile_id']}  |  "
            f"Document: {session.get('filename', 'N/A')}  |  "
            f"Date: {session.get('created_at', 'N/A')}"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(83, 96, 105)

        # Score
        score_para = doc.add_paragraph()
        score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        score_run = score_para.add_run(f"{session.get('score', 0)}/100")
        score_run.font.size = Pt(36)
        score_run.bold = True
        score = session.get("score", 0)
        score_run.font.color.rgb = RGBColor(0, 180, 0) if score >= 80 else (
            RGBColor(200, 150, 0) if score >= 50 else RGBColor(200, 0, 0)
        )

        # Summary
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(session.get("summary", ""))

        # Category scores
        cat_scores = session.get("category_scores", {})
        if cat_scores:
            doc.add_heading("Category Scores", level=1)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Light Shading Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Category"
            hdr[1].text = "Score"
            for cat, sc in cat_scores.items():
                row = table.add_row().cells
                row[0].text = cat.capitalize()
                row[1].text = str(sc)

        # Issues
        doc.add_heading(f"Issues ({len(issues)})", level=1)
        if not issues:
            doc.add_paragraph("No issues found.")
        else:
            for issue in issues:
                sev = issue.get("severity", "low")
                sev_text = {"high": "🔴 HIGH", "medium": "🟡 MEDIUM", "low": "🟢 LOW"}.get(sev, sev.upper())
                p = doc.add_paragraph()
                run = p.add_run(f"{sev_text} — {issue['message']}")
                run.bold = True

                doc.add_paragraph(
                    f"Category: {issue.get('category', '')}  |  "
                    f"Rule: {issue.get('rule_id', '')}  |  "
                    f"Confidence: {issue.get('confidence', 0)}%  |  "
                    f"Status: {issue.get('status', 'open')}",
                    style="List Bullet",
                )
                if issue.get("evidence_excerpt"):
                    doc.add_paragraph(f'Evidence: "{issue["evidence_excerpt"][:250]}"', style="Intense Quote")
                doc.add_paragraph(f"Recommendation: {issue.get('recommendation', '')}")

        # Footer
        doc.add_paragraph("")
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(
            f"Generated by ReviewMind on {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(150, 150, 150)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_content = buffer.getvalue()
        buffer.close()

        return Response(
            content=docx_content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="review_{session_id[:8]}.docx"'},
        )
    except ImportError:
        raise HTTPException(status_code=501, detail="DOCX export requires python-docx library.")
    except Exception as error:
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {error}")


# ═══════════════════════════════════════════════════════════════════════════════
# 6. PROFILES & KNOWLEDGE PACKS
# ═══════════════════════════════════════════════════════════════════════════════

@router.get("/profiles")
def list_profiles() -> list[dict[str, Any]]:
    """List all available review profiles with full metadata."""
    return engine.profiles.available()


@router.get("/profiles/{profile_id}")
def get_profile_detail(profile_id: str) -> dict[str, Any]:
    """Get full detail of a profile including permissions and requirements."""
    try:
        profile = engine.profiles.load(profile_id)
        packs = database.get_packs_for_profile(profile_id)
        return {
            "id": profile.id,
            "name": profile.name,
            "description": profile.description,
            "document_types": profile.document_types,
            "categories": profile.categories,
            "weights": profile.weights,
            "permissions": profile.permissions,
            "required_sections": profile.required_sections,
            "forbidden_sections": profile.forbidden_sections,
            "rubric": profile.rubric,
            "ai_focus": profile.ai_focus,
            "auto_fix_policy": profile.auto_fix_policy,
            "compatible_packs": packs,
        }
    except ValueError as error:
        raise HTTPException(status_code=404, detail=str(error)) from error


@router.get("/packs")
def list_packs() -> list[dict[str, Any]]:
    """List all available knowledge packs."""
    return database.list_packs()


@router.get("/packs/{pack_id}")
def get_pack_detail(pack_id: str) -> dict[str, Any]:
    """Get detail of a specific knowledge pack."""
    packs = database.list_packs()
    for pack in packs:
        if pack.get("id") == pack_id:
            return pack
    raise HTTPException(status_code=404, detail=f"Knowledge pack '{pack_id}' not found.")
